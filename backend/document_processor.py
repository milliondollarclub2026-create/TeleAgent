"""
Document Processing & RAG Module for TeleAgent
Handles file upload, smart chunking, embedding generation, and semantic search
"""
import os
import io
import re
import uuid
import logging
import tiktoken
from typing import List, Dict, Optional, Tuple
from datetime import datetime, timezone
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd
from openai import AsyncOpenAI

logger = logging.getLogger(__name__)

# Lazy initialization of OpenAI client
_openai_client = None

def get_openai_client():
    """Get or create OpenAI client with lazy initialization"""
    global _openai_client
    if _openai_client is None:
        _openai_client = AsyncOpenAI(api_key=os.environ.get('OPENAI_API_KEY'))
    return _openai_client

# Embedding model config
EMBEDDING_MODEL = "text-embedding-3-small"
EMBEDDING_DIMENSIONS = 1536

# Chunking config
MAX_CHUNK_TOKENS = 500
CHUNK_OVERLAP_TOKENS = 100
MIN_CHUNK_TOKENS = 50

# Initialize tokenizer for chunk size calculation
tokenizer = tiktoken.encoding_for_model("gpt-4")


def count_tokens(text: str) -> int:
    """Count tokens in text using tiktoken"""
    return len(tokenizer.encode(text))


def split_into_sentences(text: str) -> List[str]:
    """Split text into sentences while preserving structure"""
    # Handle common sentence endings while being careful with abbreviations
    sentence_pattern = r'(?<=[.!?])\s+(?=[A-ZА-ЯЎҚҒҲа-яўқғҳ])'
    sentences = re.split(sentence_pattern, text)
    return [s.strip() for s in sentences if s.strip()]


def create_chunks_with_overlap(
    text: str,
    source_info: str = "",
    max_tokens: int = MAX_CHUNK_TOKENS,
    overlap_tokens: int = CHUNK_OVERLAP_TOKENS
) -> List[Dict]:
    """
    Smart chunking that respects sentence boundaries and adds overlap for context continuity.
    Returns list of chunk dicts with text and metadata.
    """
    if not text or not text.strip():
        return []
    
    sentences = split_into_sentences(text)
    chunks = []
    current_chunk = []
    current_tokens = 0
    
    for sentence in sentences:
        sentence_tokens = count_tokens(sentence)
        
        # If single sentence exceeds max, split it further
        if sentence_tokens > max_tokens:
            # Flush current chunk first
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "text": chunk_text,
                    "token_count": count_tokens(chunk_text),
                    "source": source_info
                })
                current_chunk = []
                current_tokens = 0
            
            # Split long sentence by words
            words = sentence.split()
            temp_chunk = []
            temp_tokens = 0
            for word in words:
                word_tokens = count_tokens(word + " ")
                if temp_tokens + word_tokens > max_tokens and temp_chunk:
                    chunk_text = " ".join(temp_chunk)
                    chunks.append({
                        "text": chunk_text,
                        "token_count": count_tokens(chunk_text),
                        "source": source_info
                    })
                    # Keep overlap
                    overlap_words = temp_chunk[-(len(temp_chunk)//4):]
                    temp_chunk = overlap_words + [word]
                    temp_tokens = count_tokens(" ".join(temp_chunk))
                else:
                    temp_chunk.append(word)
                    temp_tokens += word_tokens
            
            if temp_chunk:
                current_chunk = temp_chunk
                current_tokens = temp_tokens
            continue
        
        # Check if adding sentence exceeds limit
        if current_tokens + sentence_tokens > max_tokens and current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "token_count": count_tokens(chunk_text),
                "source": source_info
            })
            
            # Create overlap from last sentences
            overlap_chunk = []
            overlap_tokens_count = 0
            for s in reversed(current_chunk):
                s_tokens = count_tokens(s)
                if overlap_tokens_count + s_tokens <= overlap_tokens:
                    overlap_chunk.insert(0, s)
                    overlap_tokens_count += s_tokens
                else:
                    break
            
            current_chunk = overlap_chunk + [sentence]
            current_tokens = overlap_tokens_count + sentence_tokens
        else:
            current_chunk.append(sentence)
            current_tokens += sentence_tokens
    
    # Add final chunk
    if current_chunk:
        chunk_text = " ".join(current_chunk)
        token_count = count_tokens(chunk_text)
        # Always include content, even if small (for short documents)
        if token_count > 0:
            chunks.append({
                "text": chunk_text,
                "token_count": token_count,
                "source": source_info
            })
    
    # If no chunks were created but we have original text, create a single chunk
    if not chunks and text.strip():
        chunks.append({
            "text": text.strip(),
            "token_count": count_tokens(text.strip()),
            "source": source_info
        })
    
    return chunks


# ============ File Processors ============

def process_pdf(file_content: bytes, filename: str) -> List[Dict]:
    """
    Extract text from PDF with structure preservation.
    Returns chunks with section awareness.
    """
    chunks = []
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        
        full_text = []
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text("text")
            if page_text.strip():
                full_text.append(f"[Page {page_num}]\n{page_text}")
        
        doc.close()
        
        combined_text = "\n\n".join(full_text)
        chunks = create_chunks_with_overlap(combined_text, source_info=filename)
        
        logger.info(f"PDF processed: {filename}, {len(chunks)} chunks created")
        
    except Exception as e:
        logger.error(f"PDF processing error: {e}")
        raise ValueError(f"Failed to process PDF: {str(e)}")
    
    return chunks


def process_docx(file_content: bytes, filename: str) -> List[Dict]:
    """
    Extract text from DOCX with paragraph and table handling.
    """
    chunks = []
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        
        sections = []
        current_section = {"heading": "", "content": []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"heading": text, "content": []}
            else:
                current_section["content"].append(text)
        
        # Add last section
        if current_section["content"]:
            sections.append(current_section)
        
        # Process tables
        for table in doc.tables:
            table_text = []
            headers = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    headers = cells
                else:
                    if headers:
                        row_text = ", ".join([f"{h}: {c}" for h, c in zip(headers, cells) if c])
                    else:
                        row_text = ", ".join([c for c in cells if c])
                    if row_text:
                        table_text.append(row_text)
            
            if table_text:
                sections.append({
                    "heading": "Table Data",
                    "content": table_text
                })
        
        # Create chunks from sections
        for section in sections:
            section_text = ""
            if section["heading"]:
                section_text = f"## {section['heading']}\n"
            section_text += "\n".join(section["content"])
            
            section_chunks = create_chunks_with_overlap(section_text, source_info=f"{filename} - {section['heading'] or 'Content'}")
            chunks.extend(section_chunks)
        
        logger.info(f"DOCX processed: {filename}, {len(chunks)} chunks created")
        
    except Exception as e:
        logger.error(f"DOCX processing error: {e}")
        raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    return chunks


def process_excel(file_content: bytes, filename: str) -> List[Dict]:
    """
    Process Excel/CSV files - great for product catalogs.
    Each row becomes a searchable chunk with column context.
    """
    chunks = []
    try:
        # Try to read as Excel first, fall back to CSV
        try:
            df = pd.read_excel(io.BytesIO(file_content))
        except:
            df = pd.read_csv(io.BytesIO(file_content))
        
        # Clean column names
        df.columns = [str(col).strip() for col in df.columns]
        
        # Group rows into chunks (5-10 products per chunk for context)
        rows_per_chunk = 5
        
        for i in range(0, len(df), rows_per_chunk):
            batch = df.iloc[i:i+rows_per_chunk]
            chunk_lines = []
            
            for _, row in batch.iterrows():
                # Create readable row representation
                row_parts = []
                for col, val in row.items():
                    if pd.notna(val) and str(val).strip():
                        row_parts.append(f"{col}: {val}")
                if row_parts:
                    chunk_lines.append(" | ".join(row_parts))
            
            if chunk_lines:
                chunk_text = "\n".join(chunk_lines)
                chunks.append({
                    "text": chunk_text,
                    "token_count": count_tokens(chunk_text),
                    "source": f"{filename} (rows {i+1}-{min(i+rows_per_chunk, len(df))})"
                })
        
        logger.info(f"Excel processed: {filename}, {len(chunks)} chunks created from {len(df)} rows")
        
    except Exception as e:
        logger.error(f"Excel processing error: {e}")
        raise ValueError(f"Failed to process Excel: {str(e)}")
    
    return chunks


async def process_image(file_content: bytes, filename: str) -> List[Dict]:
    """
    Use GPT-4V to describe image content for searchability.
    Great for product photos.
    """
    import base64
    
    chunks = []
    try:
        # Encode image to base64
        base64_image = base64.b64encode(file_content).decode('utf-8')
        
        # Determine image type
        if filename.lower().endswith('.png'):
            media_type = "image/png"
        elif filename.lower().endswith('.gif'):
            media_type = "image/gif"
        elif filename.lower().endswith('.webp'):
            media_type = "image/webp"
        else:
            media_type = "image/jpeg"
        
        # Call GPT-4V for description
        response = await get_openai_client().chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": """Describe this image in detail for a product knowledge base. Include:
1. What the product/item is
2. Key features visible (color, size, material, design)
3. Any text, labels, or branding visible
4. Price if shown
5. Any other relevant details for a sales agent

Format as structured text that can be searched."""
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            max_tokens=500
        )
        
        description = response.choices[0].message.content
        
        chunks.append({
            "text": f"[Image: {filename}]\n{description}",
            "token_count": count_tokens(description),
            "source": f"Image: {filename}"
        })
        
        logger.info(f"Image processed: {filename}")
        
    except Exception as e:
        logger.error(f"Image processing error: {e}")
        raise ValueError(f"Failed to process image: {str(e)}")
    
    return chunks


def process_text(content: str, title: str) -> List[Dict]:
    """Process plain text content"""
    return create_chunks_with_overlap(content, source_info=title)


# ============ Embedding Generation ============

async def generate_embedding(text: str) -> List[float]:
    """Generate embedding for a single text using OpenAI"""
    try:
        response = await get_openai_client().embeddings.create(
            model=EMBEDDING_MODEL,
            input=text,
            dimensions=EMBEDDING_DIMENSIONS
        )
        return response.data[0].embedding
    except Exception as e:
        logger.error(f"Embedding generation error: {e}")
        raise


async def generate_embeddings_batch(texts: List[str]) -> List[List[float]]:
    """Generate embeddings for multiple texts in batch"""
    try:
        response = await get_openai_client().embeddings.create(
            model=EMBEDDING_MODEL,
            input=texts,
            dimensions=EMBEDDING_DIMENSIONS
        )
        return [item.embedding for item in response.data]
    except Exception as e:
        logger.error(f"Batch embedding error: {e}")
        raise


# ============ Main Processing Function ============

async def process_document(
    file_content: bytes,
    filename: str,
    content_type: str
) -> Tuple[List[Dict], List[List[float]]]:
    """
    Main entry point for document processing.
    Returns (chunks, embeddings) tuple.
    """
    filename_lower = filename.lower()
    
    # Determine file type and process
    if filename_lower.endswith('.pdf') or content_type == 'application/pdf':
        chunks = process_pdf(file_content, filename)
    
    elif filename_lower.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        chunks = process_docx(file_content, filename)
    
    elif filename_lower.endswith(('.xlsx', '.xls')) or 'spreadsheet' in content_type or 'excel' in content_type:
        chunks = process_excel(file_content, filename)
    
    elif filename_lower.endswith('.csv') or content_type == 'text/csv':
        chunks = process_excel(file_content, filename)
    
    elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')) or content_type.startswith('image/'):
        chunks = await process_image(file_content, filename)
    
    elif filename_lower.endswith('.txt') or content_type == 'text/plain':
        text_content = file_content.decode('utf-8', errors='ignore')
        chunks = process_text(text_content, filename)
    
    else:
        # Try to process as text
        try:
            text_content = file_content.decode('utf-8', errors='ignore')
            chunks = process_text(text_content, filename)
        except:
            raise ValueError(f"Unsupported file type: {filename}")
    
    if not chunks:
        raise ValueError("No content could be extracted from the file")
    
    # Generate embeddings for all chunks
    chunk_texts = [chunk["text"] for chunk in chunks]
    embeddings = await generate_embeddings_batch(chunk_texts)
    
    return chunks, embeddings


# ============ Semantic Search ============

def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Calculate cosine similarity between two vectors"""
    import math
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    magnitude1 = math.sqrt(sum(a * a for a in vec1))
    magnitude2 = math.sqrt(sum(b * b for b in vec2))
    if magnitude1 * magnitude2 == 0:
        return 0
    return dot_product / (magnitude1 * magnitude2)


async def semantic_search(
    query: str,
    document_chunks: List[Dict],
    top_k: int = 5,
    min_similarity: float = 0.3
) -> List[Dict]:
    """
    Perform semantic search on document chunks.
    Returns top-k most relevant chunks with similarity scores.
    """
    if not document_chunks:
        return []
    
    # Generate query embedding
    query_embedding = await generate_embedding(query)
    
    # Calculate similarities
    results = []
    for chunk in document_chunks:
        if "embedding" in chunk:
            similarity = cosine_similarity(query_embedding, chunk["embedding"])
            if similarity >= min_similarity:
                results.append({
                    "text": chunk["text"],
                    "source": chunk.get("source", ""),
                    "similarity": similarity
                })
    
    # Sort by similarity and return top-k
    results.sort(key=lambda x: x["similarity"], reverse=True)
    return results[:top_k]
